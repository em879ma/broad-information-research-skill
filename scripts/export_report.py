#!/usr/bin/env python3
"""
export_report.py - Command-line tool to export research reports in various formats.

Usage:
    # As command-line tool
    python export_report.py --input results.md --format pdf
    python export_report.py --input results.md --format docx
    python export_report.py --input results.json --format csv

    # Import as module
    from export_report import export_file
    export_file(input_path, format='pdf')
"""

import argparse
import json
import os
import subprocess
import sys
from pathlib import Path
from typing import Optional, List


def get_available_formats() -> List[str]:
    """Get list of available export formats."""
    return ['md', 'txt', 'json', 'csv', 'docx', 'pdf', 'html']


def check_dependencies() -> dict:
    """Check which export dependencies are available."""
    deps = {
        'python-docx': False,
        'weasyprint': False,
        'markdown': False,
        'pdfkit': False,
        'jinja2': False,
    }
    
    try:
        import docx
        deps['python-docx'] = True
    except ImportError:
        pass
    
    try:
        import weasyprint
        deps['weasyprint'] = True
    except ImportError:
        pass
    
    try:
        import markdown
        deps['markdown'] = True
    except ImportError:
        pass
    
    try:
        import pdfkit
        deps['pdfkit'] = True
    except ImportError:
        pass
    
    try:
        import jinja2
        deps['jinja2'] = True
    except ImportError:
        pass
    
    return deps


def export_to_docx(input_path: str, output_path: Optional[str] = None) -> str:
    """Export Markdown/Text to Word (.docx) format."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Read input content
    input_path = Path(input_path)
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(input_path.stem.replace('_', ' ').title(), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Parse content and add to document
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Numbered lists
        elif line[0:2].isdigit() and '. ' in line:
            idx = line.index('. ')
            doc.add_paragraph(line[idx+2:], style='List Number')
        # Regular paragraph
        else:
            doc.add_paragraph(line)
    
    # Save document
    if output_path is None:
        output_path = str(input_path.with_suffix('.docx'))
    
    doc.save(output_path)
    return output_path


def export_to_pdf(input_path: str, output_path: Optional[str] = None) -> str:
    """Export to PDF using weasyprint or pdfkit."""
    try:
        import weasyprint
    except ImportError:
        print("Error: weasyprint not installed. Install with: pip install weasyprint", file=sys.stderr)
        raise ImportError("weasyprint not available")
    
    input_path = Path(input_path)
    
    # Read input
    with open(input_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Convert Markdown to HTML if needed
    if input_path.suffix == '.md':
        # Simple markdown to HTML conversion
        html_lines = []
        for line in html_content.split('\n'):
            line = line.strip()
            
            if not line:
                html_lines.append('<br>')
            elif line.startswith('# '):
                html_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                html_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                html_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('- ') or line.startswith('* '):
                html_lines.append(f'<li>{line[2:]}</li>')
            else:
                html_lines.append(f'<p>{line}</p>')
        
        html_content = '''
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; font-size: 12pt; }
        h1 { color: #2c3e50; }
        h2 { color: #34495e; }
        h3 { color: #7f8c8d; }
        li { margin-left: 20px; }
    </style>
</head>
<body>
''' + '\n'.join(html_lines) + '''
</body>
</html>
'''
    
    # Convert to PDF
    if output_path is None:
        output_path = str(input_path.with_suffix('.pdf'))
    
    pdf = weasyprint.HTML(string=html_content).write_pdf()
    
    with open(output_path, 'wb') as f:
        f.write(pdf)
    
    return output_path


def export_to_html(input_path: str, output_path: Optional[str] = None) -> str:
    """Export Markdown to HTML."""
    try:
        import markdown
    except ImportError:
        print("Error: markdown not installed. Install with: pip install markdown", file=sys.stderr)
        raise ImportError("markdown not available")
    
    input_path = Path(input_path)
    
    # Read input
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert to HTML
    html_content = markdown.markdown(md_content)
    
    # Wrap in basic HTML template
    full_html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{input_path.stem}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; max-width: 800px; line-height: 1.6; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; }}
        h3 {{ color: #7f8c8d; }}
        code {{ background: #f8f9fa; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background: #f8f9fa; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
        th {{ background: #3498db; color: white; }}
        tr:nth-child(even) {{ background: #f2f2f2; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
'''
    
    # Save
    if output_path is None:
        output_path = str(input_path.with_suffix('.html'))
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_html)
    
    return output_path


def export_file(input_path: str, format: Optional[str] = None, output_path: Optional[str] = None) -> str:
    """
    Export file to specified format.
    
    Args:
        input_path: Path to input file
        format: Target format ('md', 'txt', 'json', 'csv', 'docx', 'pdf', 'html')
        output_path: Optional output path (auto-generated if not provided)
    
    Returns:
        Path to exported file
    """
    input_path = Path(input_path)
    
    if format is None:
        format = input_path.suffix[1:]
    
    format = format.lower()
    
    if format == 'docx':
        return export_to_docx(str(input_path), output_path)
    elif format == 'pdf':
        return export_to_pdf(str(input_path), output_path)
    elif format == 'html':
        return export_to_html(str(input_path), output_path)
    elif format in ['md', 'txt', 'json', 'csv']:
        # Just copy the file
        if output_path is None:
            output_path = str(input_path.with_suffix(f'.{format}'))
        
        import shutil
        shutil.copy(input_path, output_path)
        return output_path
    else:
        raise ValueError(f"Unsupported format: {format}")


def main():
    parser = argparse.ArgumentParser(description='Export research reports to various formats.')
    parser.add_argument('--input', '-i', required=True, help='Input file path')
    parser.add_argument('--format', '-f', choices=get_available_formats(), help='Output format')
    parser.add_argument('--output', '-o', help='Output file path')
    parser.add_argument('--list-formats', action='store_true', help='List available formats')
    
    args = parser.parse_args()
    
    if args.list_formats:
        print("Available formats:")
        for fmt in get_available_formats():
            print(f"  - {fmt}")
        
        print("\nAvailable dependencies:")
        deps = check_dependencies()
        for dep, available in deps.items():
            status = "✓" if available else "✗"
            print(f"  {status} {dep}")
        return
    
    if not args.format:
        # Infer from input file
        args.format = Path(args.input).suffix[1:]
    
    try:
        output_path = export_file(args.input, args.format, args.output)
        print(f"Exported to: {output_path}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()